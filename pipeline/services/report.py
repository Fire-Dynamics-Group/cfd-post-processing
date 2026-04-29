"""Report orchestrator (PR 2 step 6 — real pipeline).

This module is the sidecar's only allowed entry-point into the vendored
CFD pipeline. It does NOT import legacy ``auto_report.py`` or ``main.py``
(both PySimpleGUI-coupled). The flow mirrors what those files did, but:

- No PSG popups; structured warnings + ``raise`` for fatal conditions.
- No ``is_test`` Easter egg (which hardcoded an Evelyn Court Dropbox path).
- No ``os.startfile`` auto-open — the frontend renders an explicit "Open
  in Word" button for that.
- Each phase emits ``on_progress`` so the polling UI has something to
  render. The orchestrator is CPU-bound, so the FastAPI handler runs it
  in a daemon thread (see ``pipeline/server.py``).

Public surface:

- ``ReportRequest``: pydantic model for the form payload.
- ``ReportResult``: dataclass returned on success.
- ``run_orchestrator(req, on_progress)``: the entry-point.
- ``build_context(req, scenarios_object, scenario_names, FSA, MOE,
  fds_version)``: pure-ish ctx-dict builder, exposed so the golden
  characterization tests (PR 2 step 7) can pin its shape.
"""
from __future__ import annotations

import datetime
import logging
import os
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable, Literal

# Pin matplotlib's headless backend before any chart code imports it.
# The pipeline writes PNGs to disk and never displays them; without this,
# a bundled .exe on a machine without Tk/Qt may try to load a GUI backend
# and fail. setdefault only mutates os.environ — it does not import
# matplotlib, so the lazy-import contract still holds.
os.environ.setdefault("MPLBACKEND", "Agg")

from pydantic import BaseModel, Field

from .job import Step

# NOTE on lazy imports (PR 3 decision #17):
# Heavy imports — docx, docxtpl, matplotlib (via hrr_graph), fdsreader, lxml,
# PIL (via report_draw), reportlab — are deferred into the function bodies
# that use them so that ``import pipeline.services.report`` stays cheap
# (~ms vs ~seconds). The bundled sidecar's cold start fold the first-import
# cost into the first job's run, where the user already expects a wait.

logger = logging.getLogger(__name__)


class PipelineError(Exception):
    """Expected, user-facing failure inside the orchestrator.

    Raised for conditions like "no scenarios found in PATH" or "missing
    devc.csv for scenario X" — situations where the user can fix their
    inputs and retry. The server reports these with ``ErrorType.PIPELINE``
    so the UI shows the message and step but does not surface a "Copy
    diagnostic" button (those are for InternalErrors / bugs).
    """


class ReportRequest(BaseModel):
    """Form payload — PR 2 keeps legacy parity (7 fields) plus optional OUTPUT_DIR."""

    PATH: str = Field(..., description="Path to runs' root directory")
    CLIENT_NAME: str
    PROJECT_NAME: str
    PROJECT_LOCATION: str
    EMAIL_PREFIX: str = Field(..., description="Senior's email prefix")
    HAS_EXTENDED_TRAVEL: bool = True
    MAX_TD: float | None = Field(default=None, description="Max travel distance in metres")
    GUIDANCE: Literal["BS9991", "ADB"] = "BS9991"
    OUTPUT_DIR: str | None = Field(
        default=None,
        description="Optional override for where the .docx is written. Defaults to PATH.",
    )


@dataclass
class ReportResult:
    output_path: str
    warnings: list[str]


ProgressCallback = Callable[[Step, float], None]


# ---------------------------------------------------------------------------
# Public orchestrator
# ---------------------------------------------------------------------------


def run_orchestrator(req: ReportRequest, on_progress: ProgressCallback) -> ReportResult:
    """End-to-end report generation against a real FDS run directory."""
    from docxtpl import DocxTemplate

    from scenarios_object import create_scenario_object
    from scen_object_helper_functions import return_fds_version
    from report_draw import run_all_report_draw

    warnings: list[str] = []

    # --- PARSING ----------------------------------------------------------
    on_progress(Step.PARSING, 0.0)
    if not Path(req.PATH).exists():
        raise FileNotFoundError(f"Runs directory does not exist: {req.PATH}")

    scenarios_object, scenario_names, FSA_scenarios, MoE_scenarios, error_list = (
        create_scenario_object(path_to_directory=req.PATH)
    )
    if error_list:
        # Treat errors from the parser as fatal — they indicate missing fds /
        # devc / hrr files which would cause downstream chart generation to
        # crash anyway. User-fixable, so PipelineError (not InternalError).
        raise PipelineError(
            "Scenario parse errors:\n" + "\n\n".join(error_list)
        )
    if not scenario_names:
        raise PipelineError(f"No scenarios found in {req.PATH}")

    # Resolve fds_version from the first scenario (all scenarios in a job
    # share an FDS version per legacy assumption).
    runs_root = req.PATH
    if not [f.name for f in os.scandir(runs_root) if f.is_dir()]:
        runs_root = os.path.dirname(runs_root)
    fds_version = return_fds_version(
        path_to_directory=f"{runs_root}/{scenario_names[0]}"
    )
    on_progress(Step.PARSING, 1.0)

    # --- CHARTING ---------------------------------------------------------
    on_progress(Step.CHARTING, 0.0)
    charts_dir = _resolve_charts_dir(req)
    charts_dir.mkdir(parents=True, exist_ok=True)
    _run_cfd_charts_with_progress(
        runs_root, scenario_names, str(charts_dir), on_progress
    )
    on_progress(Step.CHARTING, 1.0)

    # --- DRAWING ----------------------------------------------------------
    on_progress(Step.DRAWING, 0.0)
    template_path = _resolve_template_path()
    doc = DocxTemplate(str(template_path))
    fds_figures = run_all_report_draw(doc, runs_root, scenario_names)
    on_progress(Step.DRAWING, 1.0)

    # --- RENDERING --------------------------------------------------------
    on_progress(Step.RENDERING, 0.0)
    ctx, ref_order = build_context(
        req=req,
        scenarios_object=scenarios_object,
        scenario_names=scenario_names,
        FSA_scenarios=FSA_scenarios,
        MoE_scenarios=MoE_scenarios,
        fds_version=fds_version,
    )
    _attach_chart_images(
        ctx, doc, charts_dir, scenario_names, FSA_scenarios, MoE_scenarios
    )
    ctx["FDS_FIGURES"] = fds_figures

    output_path = _resolve_output_path(req)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.render(ctx)
    doc.save(str(output_path))
    on_progress(Step.RENDERING, 1.0)

    # --- SAVING (post-process docx tables + ref list) --------------------
    on_progress(Step.SAVING, 0.0)
    _postprocess_docx(
        output_path,
        scenarios_object=scenarios_object,
        scenario_names=scenario_names,
        FSA_scenarios=FSA_scenarios,
        MoE_scenarios=MoE_scenarios,
        ref_order=ref_order,
        fds_version=fds_version,
    )
    on_progress(Step.SAVING, 1.0)

    return ReportResult(output_path=str(output_path), warnings=warnings)


# ---------------------------------------------------------------------------
# Path resolution
# ---------------------------------------------------------------------------


def _resolve_output_dir(req: ReportRequest) -> Path:
    return Path(req.OUTPUT_DIR or req.PATH)


def _resolve_charts_dir(req: ReportRequest) -> Path:
    """Where intermediate chart PNGs are written.

    Co-located with the docx so the user can find them, but in a clearly
    named subfolder so they don't pollute OUTPUT_DIR's root.
    """
    return _resolve_output_dir(req) / f"{req.PROJECT_NAME}_charts"


def _resolve_output_path(req: ReportRequest) -> Path:
    today_str = datetime.date.today().isoformat()
    return _resolve_output_dir(req) / f"{req.PROJECT_NAME}-CFD Report-{today_str}.docx"


def _resolve_template_path() -> Path:
    return _resolve_resource("Template CFD Report.docx")


def _resolve_resource(filename: str) -> Path:
    """Find a bundled resource by walking the standard tiers.

    Order, first match wins:
      1. ``sys._MEIPASS`` — set by the PyInstaller bootloader to the dir
         where data files are extracted (onedir: ``_internal/``; onefile:
         per-run ``_MEIxxxx``).
      2. ``Path(sys.executable).parent`` — sibling of the bundled exe;
         redundant in our onedir layout but cheap to check and matches
         what users see when they unzip the install.
      3. ``Path(__file__).resolve().parents[2]`` — project root in dev.
    """
    candidates: list[Path] = []
    meipass = getattr(sys, "_MEIPASS", None)
    if meipass:
        candidates.append(Path(meipass) / filename)
    candidates.append(Path(sys.executable).parent / filename)
    candidates.append(Path(__file__).resolve().parents[2] / filename)
    for c in candidates:
        if c.exists():
            return c
    raise FileNotFoundError(
        f"{filename!r} not found. Looked in: "
        + ", ".join(str(c) for c in candidates)
    )


# ---------------------------------------------------------------------------
# Phase: charting (per-scenario progress)
# ---------------------------------------------------------------------------


def _run_cfd_charts_with_progress(
    runs_root: str,
    scenario_names: list[str],
    new_dir_path: str,
    on_progress: ProgressCallback,
) -> None:
    """Run the legacy chart pipeline scenario-by-scenario so the UI can show
    incremental progress during this phase (it dominates wall time)."""
    from helper_functions import return_paths_to_files
    from hrr_graph import run_devc_charts, run_hrr_charts

    n = len(scenario_names)
    for index, scenario_name in enumerate(scenario_names):
        path_to_hrr_file, _, path_to_fds_file, path_to_devc_file, error_list = (
            return_paths_to_files(
                scenario_name,
                dir_path=runs_root,
                new_folder_structure=True,
            )
        )
        if error_list:
            # Should have been caught by create_scenario_object; bail out
            # rather than producing a partial report.
            raise PipelineError(
                f"Missing files for scenario {scenario_name!r}: "
                + "; ".join(error_list)
            )

        firefighting = "FSA" in scenario_name
        run_hrr_charts(
            path_to_fds_file,
            path_to_hrr_file,
            new_dir_path=new_dir_path,
            firefighting=firefighting,
        )
        run_devc_charts(
            path_to_file=path_to_devc_file,
            path_to_fds_file=path_to_fds_file,
            new_dir_path=new_dir_path,
            firefighting=firefighting,
        )
        on_progress(Step.CHARTING, (index + 1) / n)


# ---------------------------------------------------------------------------
# Phase: building the docxtpl context
# ---------------------------------------------------------------------------


_NUM_TO_TEXT = {
    1: "one", 2: "two", 3: "three", 4: "four", 5: "five",
    6: "six", 7: "seven", 8: "eight", 9: "nine",
}


def _num_to_text(num: int, capitalise: bool = False) -> int | str:
    """Mirror legacy `num_to_text`: numbers 1–9 become words, 0 stays
    numeric (legacy fall-through behaviour for the no-MoE case), and
    10+ stays numeric."""
    word = _NUM_TO_TEXT.get(num)
    if word is None:
        return num
    return word.capitalize() if capitalise else word


def _compute_fire_scen_text(scenario_names: list[str], MoE_scenarios: list[str], FSA_scenarios: list[str]) -> str:
    n_total = len(scenario_names)
    text = f"{_num_to_text(n_total, capitalise=True)} fire scenario"
    if n_total > 1:
        text += "s have"
    else:
        text += " has"
    text += " been considered in this assessment"
    if not MoE_scenarios:
        text += " and will relate to the Fire Service Access phase only. The model"
        if n_total > 1:
            text += "s"
        text += " will consider the likelihood of smoke penetrating into the stair based on "
        text += "credible worst case apartment locations." if n_total > 1 else "a credible worst case apartment location."
    else:
        text += f", {_num_to_text(len(MoE_scenarios))} Means of Escape scenario"
        if len(MoE_scenarios) > 1:
            text += "s"
        text += f" and {_num_to_text(len(FSA_scenarios))} Fire Service Access scenario"
        if len(FSA_scenarios) > 1:
            text += "s"
        text += ". Fire scenarios are based on credible worst case apartment locations."
    return text


def _ref_order_for(guidance: Literal["BS9991", "ADB"]) -> list[str]:
    """Reproduce legacy `add_refs_in_order` ordering exactly."""
    order: list[str] = ["SCA_1", "BRegs"]
    if guidance == "BS9991":
        order.append("BS9991")
    else:
        order.append("ADB")
    order += ["BS7974", "FDS", "NIST"]
    if guidance == "ADB":
        order.append("BS9991")
    order += [
        "BS1366_2", "BS5839_1", "PD7974_6", "BS12101_6", "SPFE",
        "SCA_2", "PD7974_1", "BS9251", "BRE_1",
    ]
    return order


def build_context(
    req: ReportRequest,
    scenarios_object: dict[str, Any],
    scenario_names: list[str],
    FSA_scenarios: list[str],
    MoE_scenarios: list[str],
    fds_version: str,
) -> tuple[dict[str, Any], list[str]]:
    """Build the docxtpl context dict.

    Returns ``(ctx, ref_order)``. ``ref_order`` is needed by the docx
    post-processing step but is *not* a docxtpl tag; keeping it out of the
    ctx avoids leaking implementation detail into the rendered template.
    """
    from helper_functions import round_to
    from report_gen_helper_functions import scen_results_values
    from variable_text import Extended_travel_1, Extended_travel_2

    today = datetime.datetime.today()
    ctx: dict[str, Any] = {
        "PATH": req.PATH,
        "CLIENT_NAME": req.CLIENT_NAME,
        "PROJECT_NAME": req.PROJECT_NAME,
        "PROJECT_LOCATION": req.PROJECT_LOCATION,
        "EMAIL_PREFIX": req.EMAIL_PREFIX,
        "HAS_EXTENDED_TRAVEL": req.HAS_EXTENDED_TRAVEL,
        "MAX_TD": req.MAX_TD,
        # Legacy templates branch on these booleans — keep both forms.
        "BS9991": req.GUIDANCE == "BS9991",
        "ADB": req.GUIDANCE == "ADB",
        "TODAYS_DATE": today.strftime("%d-%m-%Y"),
        "scenario_names": scenario_names,
        "FDS_VERSION": fds_version,
        "HAS_SPRINKLERS": scenarios_object[scenario_names[0]]["is_sprinklered"],
    }

    fsa_scenario_nums = [i + 1 for i, s in enumerate(scenario_names) if "FSA" in s]
    moe_scenario_nums = [i + 1 for i, s in enumerate(scenario_names) if "FSA" not in s]
    ctx["fsa_nums"] = ",".join(str(e) for e in fsa_scenario_nums)
    ctx["moe_nums"] = ",".join(str(e) for e in moe_scenario_nums)

    for jinja_name, count in (
        ("NUM_SCENARIOS", len(scenario_names)),
        ("NUM_MOE_SCENARIOS", len(MoE_scenarios)),
        ("NUM_FSA_SCENARIOS", len(FSA_scenarios)),
    ):
        ctx[jinja_name] = count
        ctx[f"{jinja_name}_TEXT"] = _num_to_text(count)

    ctx["FIRE_SCEN_TEXT"] = _compute_fire_scen_text(scenario_names, MoE_scenarios, FSA_scenarios)
    ctx["FIRE_SCEN_SUB_TEXT"] = "s are" if len(scenario_names) > 1 else " is"

    if MoE_scenarios:
        ctx["MOE_SCENARIO"] = True
    if FSA_scenarios:
        ctx["FSA_SCENARIO"] = True

    if len(MoE_scenarios) == 1:
        ctx["SINGLE_MOE_SCENARIO"] = True
        scenario = MoE_scenarios[0]
        tenable_time, max_pressure_drop, _ = scen_results_values(
            scenario, scenarios_object, firefighting=False
        )
        ctx["MOE_TENABLE_TIME"] = round_to(tenable_time)
        ctx["MOE_MIN_PRESSURE"] = round_to(max_pressure_drop)
    elif len(MoE_scenarios) > 1:
        ctx["MULTIPLE_MOE_SCENARIOS"] = True
        ctx["moe_plural"] = "s"

    if len(FSA_scenarios) == 1:
        ctx["SINGLE_FSA_SCENARIO"] = True
        scenario = FSA_scenarios[0]
        text_list, worst_temp, worst_vis, _, max_pressure_drop = scen_results_values(
            scenario, scenarios_object, firefighting=True
        )
        ctx["FSA_2M_TEMP"] = text_list[0]
        if len(text_list) > 1:
            ctx["HAS_FSA_4M_TEMP"] = True
            ctx["FSA_4M_TEMP"] = text_list[1]
            if len(text_list) > 2:
                ctx["HAS_FSA_15M_TEMP"] = True
                ctx["FSA_15M_TEMP"] = text_list[2]
        ctx["FSA_MIN_PRESSURE"] = max_pressure_drop
        ctx["FSA_STAIR_VIS"] = round_to(worst_vis)
        ctx["FSA_STAIR_TEMP"] = round_to(worst_temp)
    elif len(FSA_scenarios) > 1:
        ctx["MULTIPLE_FSA_SCENARIOS"] = True
        ctx["fsa_plural"] = "s"

    if req.HAS_EXTENDED_TRAVEL:
        ctx["EXTENDED_TD_1"] = Extended_travel_1
        ctx["EXTENDED_TD_2"] = Extended_travel_2

    ref_order = _ref_order_for(req.GUIDANCE)
    for ref in ref_order:
        ctx[f"REF_{ref}"] = ref_order.index(ref) + 1

    return ctx, ref_order


# ---------------------------------------------------------------------------
# Phase: chart-image attachment to the docxtpl context
# ---------------------------------------------------------------------------


def _attach_chart_images(
    ctx: dict[str, Any],
    doc: "DocxTemplate",
    charts_dir: Path,
    scenario_names: list[str],
    FSA_scenarios: list[str],
    MoE_scenarios: list[str],
) -> None:
    """Load each PNG produced by the charting phase as an InlineImage and
    attach it to the ctx under the keys the docxtpl template expects.

    Mirrors the legacy ``insert_charts`` (main.py:273-313).
    """
    from docx.shared import Inches
    from docxtpl import InlineImage
    from helper_functions import find_all_files_of_type, group_charts_by_scenario

    chart_names = find_all_files_of_type(str(charts_dir), suffix=".png")
    charts_by_scenario = group_charts_by_scenario(chart_names, scenario_names)

    inline: dict[str, InlineImage] = {}
    for chart in chart_names:
        inline[chart] = InlineImage(
            doc,
            image_descriptor=str(charts_dir / chart),
            width=Inches(6),
            height=Inches(4),
        )

    appendix: list[dict[str, Any]] = []
    for i, scenario_name in enumerate(scenario_names):
        current_charts = charts_by_scenario[i]
        scen_data: dict[str, Any] = {
            "index": i + 1,
            "name": scenario_name,
            "type": "FSA" if "FSA" in scenario_name else "MOE",
        }
        for chart_type in ("hrr", "vis", "temp", "pres", "vel"):
            matches = [c for c in current_charts if chart_type in c.lower()]
            if matches:
                for c in matches:
                    suffix = "_STAIR" if "stair" in c else ""
                    key = f"SCEN_{i+1}_{chart_type.upper()}{suffix}_CHART"
                    ctx[key] = inline[c]
                    scen_data[f"{chart_type.upper()}{suffix}_CHART"] = inline[c]
            else:
                scen_data[f"{chart_type.upper()}_CHART"] = ""
        appendix.append(scen_data)
    ctx["APPENDIX"] = appendix


# ---------------------------------------------------------------------------
# Phase: post-processing the rendered docx
# ---------------------------------------------------------------------------


def _replace_table_cell_content(
    cell: Any, text: str, *, is_bold: bool = False, alignment: int = 1
) -> None:
    from docx.shared import Pt, RGBColor
    from constants import font_name_light

    cell.text = text
    paragraphs = cell.paragraphs
    paragraphs[0].alignment = alignment
    runs = paragraphs[0].runs
    if not runs:
        return
    font = runs[0].font
    font.size = Pt(9)
    font.name = font_name_light
    font.color.rgb = RGBColor(64, 64, 64)
    font.bold = is_bold


def _delete_row_in_table(table: Any, row: int) -> None:
    table._tbl.remove(table.rows[row]._tr)


def _alter_table_rows(total_rows: int, table: Any, header_rows: int = 1) -> None:
    rows_to_remove = len(table.rows) - (total_rows + header_rows)
    for _ in range(rows_to_remove):
        _delete_row_in_table(table, row=-1)


def _delete_paragraph(paragraph: Any) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def _locate_tables_by_cell(
    document: Any, row_idx: int, col_idx: int, cell_text: str
) -> list[Any]:
    matches: list[Any] = []
    for table in document.tables:
        cell = table.rows[row_idx].cells[col_idx]
        if cell_text in cell.text:
            matches.append(table)
    return matches


def _populate_scenario_table(
    table: Any,
    scenario_names: list[str],
    scenarios_object: dict[str, Any],
) -> None:
    from helper_functions import round_to

    _alter_table_rows(total_rows=len(scenario_names), table=table)
    for idx, name in enumerate(scenario_names):
        row_index = idx + 1
        row_cells = table.rows[row_index].cells
        venting_obj = scenarios_object[name]["venting"]
        mech_extract = venting_obj["mech_extract"]
        mech_supply = venting_obj["mech_supply"]
        for col in range(1, len(row_cells)):
            if col == 1:
                _replace_table_cell_content(
                    row_cells[col],
                    "Fire Service Access" if "FSA" in name else "Means of Escape",
                )
            elif col == 3:
                total_extract = round_to(mech_extract["number"] * mech_extract["flow"])
                _replace_table_cell_content(row_cells[col], str(total_extract))
            elif col == 4:
                cell_text_list: list[str] = []
                total_supply = round_to(mech_supply["number"] * mech_supply["flow"])
                if total_supply:
                    cell_text_list.append(f"Mechanical Supply – {total_supply} m3/s")
                aov = venting_obj["stair_aov"]["area"]
                if aov:
                    cell_text_list.append(f"{aov} m2 AOV")
                cell_text = "; ".join(cell_text_list)
                _replace_table_cell_content(
                    row_cells[col],
                    f"ENGINEER TO CONFIRM: {cell_text}",
                )


def _populate_timeline_table(
    table: Any,
    scenarios_object: dict[str, Any],
    primary_scenario: str,
) -> None:
    timings = scenarios_object[primary_scenario]
    door_times = timings["door_opening_times"]
    end_time = timings["end_time"]
    keymap = {
        "Apartment Door Open": door_times["opening_apartment"],
        "Stair Door Open": door_times["opening_stair"],
        "Apartment Door Close": door_times["closing_apartment"],
        "Stair Door Close": door_times["closing_stair"],
        "Terminate": end_time,
    }
    for row in table.rows:
        row_cells = row.cells
        title_text = row_cells[0].text
        for label, value in keymap.items():
            if label in title_text:
                _replace_table_cell_content(row_cells[1], str(value))
                break


def _populate_results_table(
    document: Any,
    scenarios_object: dict[str, Any],
    *,
    firefighting: bool,
    scenarios: list[str],
    scenario_nums: list[int],
) -> None:
    from helper_functions import round_to
    from report_gen_helper_functions import scen_results_values

    results_tables = _locate_tables_by_cell(document, 0, -1, "Meets Performance")
    if not firefighting:
        moe_results_table = results_tables[0]
        _alter_table_rows(total_rows=len(scenarios), table=moe_results_table)
        for idx, scenario in enumerate(scenarios):
            row_index = idx + 1
            row_cells = moe_results_table.rows[row_index].cells
            tenable_time, max_pressure_drop, meet_criteria = scen_results_values(
                scenario, scenarios_object, firefighting=False
            )
            _replace_table_cell_content(row_cells[0], str(scenario_nums[idx]))
            _replace_table_cell_content(row_cells[1], str(round_to(tenable_time)))
            _replace_table_cell_content(row_cells[2], f"{max_pressure_drop}kPa")
            _replace_table_cell_content(row_cells[3], meet_criteria)
    else:
        fsa_results_table = results_tables[-1]
        _alter_table_rows(total_rows=len(scenarios), table=fsa_results_table, header_rows=2)
        for idx, scenario in enumerate(scenarios):
            row_index = idx + 2
            row_cells = fsa_results_table.rows[row_index].cells
            _replace_table_cell_content(row_cells[0], str(scenario_nums[idx]))
            text_list, worst_temp, worst_vis, meet_criteria, max_pressure_drop = (
                scen_results_values(scenario, scenarios_object, firefighting=True)
            )
            for j, text in enumerate(text_list):
                _replace_table_cell_content(row_cells[j + 1], text)
            _replace_table_cell_content(row_cells[4], str(round_to(worst_vis)))
            _replace_table_cell_content(row_cells[5], str(round_to(worst_temp)))
            _replace_table_cell_content(row_cells[-2], str(max_pressure_drop))
            _replace_table_cell_content(row_cells[-1], meet_criteria)


def _populate_references(
    document: Any, ref_order: list[str], fds_version: str
) -> None:
    """Populate the references section by mapping ``REF_<id>`` paragraphs to
    rows in ``references.csv``.
    """
    ref_repo_file = _resolve_references_csv()
    with open(ref_repo_file, "r+", encoding="utf8") as f:
        ref_repo_list = f.readlines()[1:]
    split_repo_list = [line.split(",") for line in ref_repo_list]

    ref_table_paras = [p for p in document.paragraphs if "REF_" in p.text]

    for ref_idx, target_id in enumerate(ref_order):
        for line in split_repo_list:
            id_, ref_title, *ref_info = line
            if target_id != id_:
                continue
            ref_info_text = ",".join(str(item) for item in ref_info).strip("\n")
            para = ref_table_paras[ref_idx]
            para.alignment = 0
            para.clear()
            if id_ == "FDS":
                ref_title = f"FDS Version {fds_version}"
            run = para.add_run(ref_title.replace('"', ""))
            run.bold = True
            para.add_run(" ")
            para.add_run(ref_info_text.replace('"', ""))
            break

    # Remove unused rows (template ships with more rows than any one report uses).
    rows_to_remove = len(ref_table_paras) - len(ref_order)
    for i in range(rows_to_remove):
        _delete_paragraph(ref_table_paras[-(i + 1)])


def _resolve_references_csv() -> Path:
    return _resolve_resource("references.csv")


def _postprocess_docx(
    output_path: Path,
    *,
    scenarios_object: dict[str, Any],
    scenario_names: list[str],
    FSA_scenarios: list[str],
    MoE_scenarios: list[str],
    ref_order: list[str],
    fds_version: str,
) -> None:
    """Re-open the rendered docx with python-docx and apply table edits
    that docxtpl can't express (cell-by-cell formatting, row deletion,
    references list).
    """
    from docx import Document

    document = Document(str(output_path))

    # Scenario table is fixed at index 2 in the legacy template.
    scenario_table = document.tables[2]
    _populate_scenario_table(scenario_table, scenario_names, scenarios_object)

    timeline_tables = _locate_tables_by_cell(document, 0, 0, "Event")
    if MoE_scenarios:
        _populate_timeline_table(
            timeline_tables[0], scenarios_object, MoE_scenarios[0]
        )
    if FSA_scenarios:
        _populate_timeline_table(
            timeline_tables[-1], scenarios_object, FSA_scenarios[0]
        )

    fsa_scenario_nums = [
        i + 1 for i, s in enumerate(scenario_names) if "FSA" in s
    ]
    moe_scenario_nums = [
        i + 1 for i, s in enumerate(scenario_names) if "FSA" not in s
    ]

    if len(MoE_scenarios) > 1:
        _populate_results_table(
            document,
            scenarios_object,
            firefighting=False,
            scenarios=MoE_scenarios,
            scenario_nums=moe_scenario_nums,
        )
    if len(FSA_scenarios) > 1:
        _populate_results_table(
            document,
            scenarios_object,
            firefighting=True,
            scenarios=FSA_scenarios,
            scenario_nums=fsa_scenario_nums,
        )

    _populate_references(document, ref_order, fds_version)

    document.save(str(output_path))
