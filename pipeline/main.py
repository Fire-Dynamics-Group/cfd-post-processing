from pathlib import Path
import datetime
import sys
import os

print("path: ",os.path.dirname(sys.executable))
import PySimpleGUI as sg
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Inches, Pt, RGBColor
from docx import Document

# https://docxtpl.readthedocs.io/en/latest/
from constants import font_name_light
from variable_text import Extended_travel_1, Extended_travel_2
from helper_functions import round_to, find_all_files_of_type, group_charts_by_scenario
from scenarios_object import create_scenario_object
from scen_object_helper_functions import return_fds_version
from report_gen_helper_functions import scen_results_values
from hrr_graph import run_CFD_charts
from validate import validate_form, generate_error_message, scenario_types
from report_draw import run_all_report_draw

document_name = "Template CFD Report.docx"
# check if file exists, if not use binary
document_path = f"{document_name}" # Path(__file__).parent /"CFD Word Template"/document_name
if os.path.exists(document_path):
    doc = DocxTemplate(document_path)



def Delete_row_in_table(table, row):
    if type(table) == int:
        table_object = document.tables[table]
    else:
        table_object = table
    table_object._tbl.remove(table_object.rows[row]._tr)


''' 
    TODO: have gui popup first
    allow path to be entered by user
'''
# Delete_row_in_table(2, -1)

# print(document.tables[0])

# for tables in document.tables:
#     for rows in tables.rows:
#         for cells in rows.cells:
#             print(cells.text)

# document = Document(document_path)
# get all styles from doc
# may need to get index of paragraph and match to styles each time
# search for curly braces??
# styles = document.styles
# paragraph_styles = [
#     s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH
# ]





def create_inline_image(image_file, template=doc):
    return InlineImage(template, image_descriptor=image_file, width=Inches(6), height=Inches(4)) # can add relative path entry box??
# charts= {}
# for chart in chart_names:
#     charts[chart] = create_inline_image(f"png_charts/{chart}", template=doc)


today = datetime.datetime.today()

'''TODO: move gui to own script -> allowing path to be entered manually '''
# GUI input boxes etc in layout
layout = [
       [sg.Text("Path to runs' root directory:"), sg.Input(key="PATH", do_not_clear=True)],        
       [sg.Text("Client name:"), sg.Input(key="CLIENT_NAME", do_not_clear=True)],
        [sg.Text("Project name:"), sg.Input(key="PROJECT_NAME", do_not_clear=True)],
        [sg.Text("Project Locations:"), sg.Input(key="PROJECT_LOCATION", do_not_clear=True, size=(20,1))],
        [sg.Text("Senior's email prefix:"), sg.Input(key="EMAIL_PREFIX", do_not_clear=True, size=(20,1))],
        [sg.Text("Extended Travel Distances:"), 
        sg.Radio('True', group_id="EXTENDED_TRAVEL", key="HAS_EXTENDED_TRAVEL",default=True),
        sg.Radio('False', group_id="EXTENDED_TRAVEL", key="NO_EXTENDED_TRAVEL")],
        # if values["HAS_EXTENDED_TRAVEL"]:
        [sg.Text("Max Travel Distance:"), sg.Input(key="MAX_TD", do_not_clear=True, size=(20,1)), sg.Text("m")],
        # guidance doc
        [sg.Text("Guidance Doc:"), 
        sg.Radio('BS9991', group_id="GUIDANCE", key="BS9991",default=True),
        sg.Radio('ADB', group_id="GUIDANCE", key="ADB")],
        # [sg.Text("Extended Travel Distances:"), sg.Listbox((["True", "False"]), size=(20,4), enable_events=False, key='HAS_EXTENDED_TRAVEL')],
        [sg.Button("Create Report"), sg.Exit()], 
]

window = sg.Window("Report Generator", layout, element_justification="right")

while True:
    event, values = window.read()
    if event == sg.WIN_CLOSED or event == "Exit":
        break
    if event == "Create Report":
        ''' 
        TODO: verify input and template file etc 
        check path
        '''
        if values['PATH']:
            values['PATH'] = r"{}".format(values['PATH'])
        is_valid, values_invalid = validate_form(values)
        files_error_message = ""
        if is_valid == False:
            error_message = generate_error_message(values_invalid)
            sg.popup_error(error_message,title="Form Input Error")
        else:
            path_to_root_directory = f"{values['PATH']}"
            # path_to_root_directory = Path(r"C:\Users\IanShaw\Dropbox\Projects CFD\9. 100 Avenue Road\Jan 2023 Corridor Models")
            # path_to_root_directory = Path(r"C:\Users\IanShaw\Dropbox\Projects CFD\22. Sweet Street\Resi\Final")
            # path_to_root_directory = Path(r"C:\Users\IanShaw\Dropbox\Projects CFD\1. Graph Generation\Test Cases\Test4")
            
            '''
            # TODO: check folder structure
            # if no subfolder -> message to user to add subfolder
            # if fds file not named appropriately -> message to user to rename
            # 
            '''
            # have path_to_dir changed if required
            # trial_path = f'{path_to_root_directory}/{scenario_name}'

            scenarios_object, scenario_names, FSA_scenarios, MoE_scenarios, error_list = create_scenario_object(path_to_directory=path_to_root_directory)
            if len(error_list) > 0:
                sg.popup_error("Error", '\n\n'.join(error_list))
            # have popup -> x MOE runs, y FSA runs -> ask to continue or rename folders with 'FSA'
            files_error_message = scenario_types(FSA_scenarios, MoE_scenarios)
            # sg.popup_error(files_error_message,title="Form Input Error")
            sg.popup("Scenarios Found:", files_error_message)
        if is_valid and len(error_list) == 0:        
            values["TODAYS_DATE"] = today.strftime("%d-%m-%Y")

            # TODO: get path from user - perhaps initial input before further one
            # C:\Users\IanShaw\Dropbox\Projects CFD\25. Claridges\Runs
            # path_to_root_directory = Path(r"C:\Users\IanShaw\Dropbox\Projects CFD\26. Breams Building\Runs")

            if len([ f.name for f in os.scandir(path_to_root_directory) if f.is_dir() ]):
                # path_to_root_directory = path_to_root_directory
                pass
            else:
                path_to_root_directory = os.path.dirname(path_to_root_directory) # need path?
            # go first scenario - all should be the same version
            fds_version = return_fds_version(path_to_directory=f'{path_to_root_directory}/{scenario_names[0]}')

            values["HAS_SPRINKLERS"] = scenarios_object[scenario_names[0]]["is_sprinklered"]

            def num_to_text(num, capatalise=False):
                if num < 10:
                    if num == 1:
                        num = 'one'
                    elif num ==2:
                        num = "two"
                    elif num == 3:
                        num = 'three'
                    elif num == 4:
                        num = 'four'
                    elif num ==5:
                        num = "five"
                    elif num == 6:
                        num = 'six'
                    elif num == 7:
                        num = 'seven'
                    elif num == 8:
                        num = 'eight'
                    elif num ==9:
                        num = "nine"
                    if capatalise == True:
                        num = num.capitalize()
                return num
            
            def jinja_num_and_text(jinja_name, number):
                values[jinja_name] = number
                values[f'{jinja_name}_TEXT'] = num_to_text(number) # later run through above function
            
            jinja_num_and_text("NUM_SCENARIOS", len(scenario_names))
            jinja_num_and_text("NUM_MOE_SCENARIOS", len(MoE_scenarios))
            jinja_num_and_text("NUM_FSA_SCENARIOS", len(FSA_scenarios))

            # TODO: move variable texts to own page
            def compute_fire_scen_text():
                first = f'{num_to_text(len(scenario_names), capatalise=True)} fire scenario'
                if len(scenario_names) > 1:
                    first += 's have'
                else:
                    first += ' has'
                first += ' been considered in this assessment'
                if len(MoE_scenarios) == 0:
                    first += ' and will relate to the Fire Service Access phase only. The model'
                    if len(scenario_names) > 1:
                        first += 's'
                    first += ' will consider the likelihood of smoke penetrating into the stair based on '
                    if len(scenario_names) > 1:
                        first += 'credible worst case apartment locations.'
                    else:
                        first += 'a credible worst case apartment location.'
                else:
                    first += f', {num_to_text(len(MoE_scenarios))} Means of Escape scenario'
                    if len(MoE_scenarios) > 1:
                        first += 's'
                    first += f' and {num_to_text(len(FSA_scenarios))} Fire Service Access scenario'
                    if len(FSA_scenarios) > 1:
                        first += 's'
                    first += '. Fire scenarios are based on credible worst case apartment locations.'
                return first
                # if len(MoE_scenarios) == 0:
                # pass

            fire_scenario_text = compute_fire_scen_text()
            values["FIRE_SCEN_TEXT"] = fire_scenario_text
            if len(scenario_names) > 1:
                fire_scenario_sub_text = 's are'
            else:
                fire_scenario_sub_text = ' is'
            values["FIRE_SCEN_SUB_TEXT"] = fire_scenario_sub_text

            # TODO: how to auto generate charts?
            values["FDS_VERSION"] = fds_version
            # also confirm naming conventions
            def ref_number(ref_order):
                for ref in ref_order:
                    values[f'REF_{ref}']= ref_order.index(ref)+1
            # TODO: pull reference order from jinja tags in document -> then remove adb/bs9991
            def add_refs_in_order():
                ref_order = ["SCA_1", "BRegs"]
                if values["BS9991"] == True:
                    ref_order.append("BS9991")
                else:
                    ref_order.append("ADB")

                    # index + 3
                    # values["REF_ADB"] = ref_order.index("ADB")

                # TODO: Future bring in ref's programattically from template
                # TODO: other refs
                ref_order.append("BS7974")
                ref_order.append("FDS")
                ref_order.append("NIST")
                if values["BS9991"] == False:
                    ref_order.append("BS9991")

                ref_order.append("BS1366_2")
                # L5 ref??
                ref_order.append("BS5839_1")
                ref_order.append("PD7974_6")
                ref_order.append("BS12101_6")
                ref_order.append("SPFE")
                ref_order.append("SCA_2")

                ref_order.append("PD7974_1")
                ref_order.append("BS9251")
                ref_order.append("BRE_1")
                return ref_order


            # jinja should output x amount of scenario loops?
            # TODO: use output path above for images to be saved to
            if os.path.isdir("outputReports"):
                new_dir_path = f"outputReports/{values['PROJECT_NAME']}" #Path(__file__).parent / "outputReports"/f"{values['PROJECT_NAME']}"
            else:
                new_dir_path = f"{values['PROJECT_NAME']}"
            
            os.mkdir(new_dir_path)
            
            run_CFD_charts(path_to_root_directory, scenario_names, new_dir_path)

            appendix_obj = []
            # scope images in folder
            def insert_charts(new_dir_path):
                # look at charts in folder
                chart_names = find_all_files_of_type(new_dir_path, suffix=".png")
                # Group charts by scenario via the FDS-file stem (handles
                # prefix collisions like FS1 vs FS10 and project-numbered jobs
                # where every chart starts with the same digits).
                charts_by_scenario = group_charts_by_scenario(chart_names, scenario_names)
                charts= {}
                for chart in chart_names:
                    charts[chart] = create_inline_image(f"{new_dir_path}/{chart}", template=doc)
                # filter for each scenario
                # input in order: hrr, vis, (vs time moe/vs distance fsa)temp, zslice, max p drop
                # then
                i = 0
                for i in range(len(scenario_names)):
                    current_scen_chart_names = charts_by_scenario[i]
                    # add scenario to appendix object
                    # add whether 4m is available if FSA corridor temp
                    current_scen_data = {"index": i+1, "name": scenario_names[i], "type": "FSA" if "FSA" in scenario_names[i] else "MOE"}
                    def set_chart(chart_type = "hrr"):
                        current_chart_list = [f for f in current_scen_chart_names if chart_type in f.lower()]
                        if len(current_chart_list) > 0:
                            for idx in range(len(current_chart_list)):
                                current_chart = current_chart_list[idx]
                            # current_chart = current_list

                                values[f"SCEN_{i+1}_{chart_type.upper()}{'_STAIR' if 'stair' in current_chart else ''}_CHART"] = charts[current_chart]
                                chart_obj = charts[current_chart]
                                current_scen_data[f"{chart_type.upper()}{'_STAIR' if 'stair' in current_chart else ''}_CHART"] = chart_obj
                        else:
                            chart_obj = ""
                            current_scen_data[f"{chart_type.upper()}_CHART"] = chart_obj

                            # add to object here
                    set_chart("hrr")
                    set_chart("vis")
                    set_chart("temp")
                    set_chart("pres")
                    set_chart("vel")
                    appendix_obj.append(current_scen_data)
                values["APPENDIX"] = appendix_obj
                print("test")

                
            # # should have scenario name/number in title of charts
            # values[f"SCENARIO_{scenario_index}_PRESSURE_CC_MOE"] = charts["Pressure cc_pres__chart.png"] 
            # values["SCENARIO_1_HRR_CC_MOE"] = charts["hrr_chart.png"]
            # values["SCENARIO_1_VELOCITY_CC_MOE"] = charts["Velocity cc_vel__chart.png"]
            # values["SCENARIO_1_VISIBILITY_CC_MOE"] = charts["Visibility cc_vis__chart.png"]
            insert_charts(new_dir_path)
            print("values: ", values)
            # TODO: scope if only one of either FSA or MOE and if so input text into bullets for scen overview using jinja
            # MOE_SCENARIO 
            if len(MoE_scenarios) > 0:
                values["MOE_SCENARIO"] = True

            if len(FSA_scenarios) > 0:
                values["FSA_SCENARIO"] = True
            # values["MOE_SCENARIO"] = True
            if len(MoE_scenarios) == 1:
                values["SINGLE_MOE_SCENARIO"] = True
                # MOE_TENABLE_TIME
                # MOE_MIN_PRESSURE
                scenario = MoE_scenarios[0]
                tenable_time, max_pressure_drop, meet_criteria = scen_results_values(scenario, scenarios_object, firefighting=False)
                values["MOE_TENABLE_TIME"] = round_to(tenable_time)
                values["MOE_MIN_PRESSURE"] = round_to(max_pressure_drop)
    # if len(FSA_scenarios) == 1:
            if len(FSA_scenarios) == 1:
                values["SINGLE_FSA_SCENARIO"] = True
                # MOE_TENABLE_TIME
                # MOE_MIN_PRESSURE
                scenario = FSA_scenarios[0]
                text_list, worst_temp, worst_vis, meet_criteria, max_pressure_drop = scen_results_values(scenario, scenarios_object, firefighting=True)
                # length of text list -> if 4m or 15m available
                values["FSA_2M_TEMP"] = text_list[0]
                # TODO: do not render if 'N/A'
                if len(text_list) > 1:
                    values["HAS_FSA_4M_TEMP"] = True
                    values["FSA_4M_TEMP"] = text_list[1]
                    if len(text_list) > 2:
                        values["HAS_FSA_15M_TEMP"] = True
                        values["FSA_15M_TEMP"] = text_list[2]  
                values["FSA_MIN_PRESSURE"] = max_pressure_drop
                values["FSA_STAIR_VIS"] = round_to(worst_vis)
                values["FSA_STAIR_TEMP"] = round_to(worst_temp)
        
            if len(MoE_scenarios) > 1:
                values["MULTIPLE_MOE_SCENARIOS"] = True
            if len(FSA_scenarios) > 1:
                values["MULTIPLE_FSA_SCENARIOS"] = True           # populate table using docx

            # perhaps only some scenarios have extended TD's?
            if values["HAS_EXTENDED_TRAVEL"] == True:
                values["EXTENDED_TD_1"] = Extended_travel_1
                values["EXTENDED_TD_2"] = Extended_travel_2
                if values["BS9991"] == True: # both TD's and 991
                    pass
                else: # both TD's and ADB
                    pass
            else: # no extended TD's
                if values["BS9991"] == True: # no TD's and 991
                    pass
                else: # no TD's and ADB -> find out td's etc
                    pass

            if values["BS9991"] == True: # below irrespective of TD's
                # inputs
            #     values["Guidance_1"] = BS9991_1
            #     values["Guidance_2"] = BS9991_2
            # else:
            #     values["Guidance_1"] = ADB_1
            #     values["Guidance_2"] = ADB_2
                pass
            # TODO: scope if sprinklers provided from fds file
            # must be after button pressed!!
            ref_order = add_refs_in_order()
            ref_number(ref_order)
            '''
                fds figure object

            '''
            # dwelling path, scenario,
            fig_charts = run_all_report_draw(doc, path_to_root_directory, scenario_names)
            values["FDS_FIGURES"] = fig_charts
            # Render the template, save new word document & inform user
            doc.render(values)
            output_path = f"{new_dir_path}/{values['PROJECT_NAME']}-CFD Report.docx"
            doc.save(output_path)
            sg.popup("File saved", f"File has been saved here: {output_path}")
            # then apply docx to tables etc
            # TODO: docx separate script, output_path sent as parameter, import object
            document = Document(output_path)

                    # loop below through each row and column
            def replace_table_cell_content(cell, replacement_text, is_bold=False, alignment=1):
                # later make object with fonts etc
                cell.text = replacement_text
                paragraphs = cell.paragraphs
                paragraphs[0].alignment = alignment # 1 = centered
                # for run in paragraphs.runs:
                run = paragraphs[0].runs
                font = run[0].font
                font.size = Pt(9) # pull from object
                font.name = font_name_light
                font.color.rgb = RGBColor(64,64,64) #gray color # to be pulled from object
                font.bold = is_bold

            # TODO: references: number superscript in report using Jinjs and 
            # TODO: have in order in references section 
            # TODO: pull references from csv
            ref_repo_file = Path(__file__).parent / 'references.csv'
            # TODO: remove unused lines
            with open(ref_repo_file, "r+", encoding="utf8") as f:
                ref_repo_list = f.readlines()[1:]
            split_repo_list = [f.split(",") for f in ref_repo_list] # id, title, ...ref_info
            document_text = [para.text for para in document.paragraphs]
            ref_table_paras = [para for para in document.paragraphs if "REF_" in para.text]

            # reference still accessible using below
            def delete_paragraph(paragraph):
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None
            
            # TODO: follow order of ref_order list
            # TODO: remove unused rows
            for ref_table_index in range(len(ref_order)):
            # ref_table_index = 0
                # TODO: find id from csv
                target_id = ref_order[ref_table_index]
                # loop through lines -> find same id 
                for csv_index in range(len(split_repo_list)):
                    test_line = split_repo_list[csv_index]
                    id, ref_title, *ref_info = test_line
                    # TODO: if fds -> insert version used from fds_version
                    if target_id == id:
                        # target_line = test_line
                        ref_info = ','.join(str(item) for item in ref_info)
                        ref_info.strip('\n')
                        para = ref_table_paras[ref_table_index]
                        para.alignment = 0
                        para.clear()
                        if id == "FDS":
                            ref_title = f'FDS Version {fds_version}'
                        run = para.add_run(ref_title.replace('"', ''))
                        run.bold = True
                        para.add_run(" ")
                        para.add_run(ref_info.replace('"', ''))
                        break
            # for index in range(10):
            #     delete_paragraph(ref_table_paras[-(index+1)])
            # TODO: remove excess rows
            def delete_excess_refs(starting_rows, required_rows, paragraphs):
                rows_to_remove = starting_rows - required_rows
                for index in range(rows_to_remove):
                    delete_paragraph(paragraphs[-(index+1)])
                # number required => len(ref_order)
                # total at beginning => ref_table_paras
            delete_excess_refs(starting_rows=len(ref_table_paras), required_rows=len(ref_order), paragraphs=ref_table_paras)
            # use length of ref_list
            
            len(ref_order)
            def find_paragraghs_containing_string(target_string):
                return [para.text for para in document.paragraphs if target_string in para.text]
            # find [REF]
            # use index or filter where "[REF]" in para.text 
            # ref_paragraphs = find_paragraghs_containing_string(target_string="[REF]")
            


            # # find reference from text, or previously mark?? with ref id
            # # central version can have id's 
            # # future pull from excel doc
            # # or object e.g. refs.adb
            # # add id's
            # # find ids next to ref
            # # loop through object keys
            # # paras_with_refs = [para.text for para in document.paragraphs if "[REF]" in para.text]
            # ref_keys = list(reference_object.keys())
            # for index, value in enumerate(ref_paragraphs):
            #     if any(x in value for x in ref_keys):
            #         ref_id_strings = [f for f in ref_paragraphs[index].split() if "REF" in f]
            #         # for each get the ref_key
            #         for id_string in ref_id_strings:
            #             for item in ref_keys:
            #                 if item in id_string:
            #                     pass
            #                     # append to ref_array 
            #                     # find length of ref_array -> number
            #                     # insert number - formatted etc

            # TODO: find reference table
            # extract references from csv
            # style appropriately
            # order from ref_order list
            # remove additional

            def alter_table_rows(total_rows, table, header_rows = 1):
                current_rows = len(table.rows)
                required_rows = total_rows + header_rows
                # remove bottom rows
                rows_to_remove = current_rows - required_rows
                # loop remove row[-1] * rows_to_remove
                for i in range(rows_to_remove):
                    Delete_row_in_table(table, row=-1)

            def is_scenario_firefighting(name):
                if 'FSA'in name:
                    return True

            def reformat_table_cell(cell):
                paragraphs = cell.paragraphs
                paragraphs[0].alignment = 1 # 1 = centered
                # for run in paragraphs.runs:
                run = paragraphs[0].runs
                font = run[0].font
                font.size = Pt(9) # pull from object
                font.name = font_name_light
                font.color.rgb = RGBColor(64,64,64) #gray color # to be pulled from object



            scenario_table = document.tables[2] # later find tables by name below??
            alter_table_rows(total_rows=len(scenario_names), table=scenario_table)
            # loop through scenarios

            # loop through each scenario in list
            def scenario_table_row(table_object, index_scenario):
                # column zero already has numbers
                model_name = scenario_names[index_scenario]
                row_index = index_scenario + 1
                row_object = table_object.rows[row_index]
                row_cells = row_object.cells
                num_cells = len(row_object.cells)

                venting_obj =scenarios_object[model_name]["venting"]
                mech_vent_obj = venting_obj["mech_extract"]
                for column_index in range(1, num_cells):
                    # column 0 should be populated with numbers
                # column 1 - type
                    if column_index == 1:
                        is_firefighting = is_scenario_firefighting(model_name)
                        if is_firefighting:
                            cell_text = 'Fire Service Access'
                        if not is_firefighting:
                            cell_text = 'Means of Escape'
                        replace_table_cell_content(cell=row_cells[column_index], replacement_text=cell_text) # how is it getting to other rows?
                # column 2 Reason for Modelling - for engineer
                # column 3 mech extract rate
                    if column_index ==3:
                        # model_name
                        total_extract = (mech_vent_obj["number"] * mech_vent_obj["flow"])

                        total_extract= round_to(value=total_extract)
                        cell_text = str(total_extract)

                        replace_table_cell_content(cell=row_cells[column_index], replacement_text=cell_text)

                    # if column_index ==4:
                    if column_index ==4:
                        cell_text_list = [] # join with semi colons -> cell_text
                        # TODO: how to superscript??
                        total_supply = round_to(venting_obj["mech_supply"]["number"] * venting_obj["mech_supply"]["flow"])
                        if total_supply:
                            cell_text_list.append(f'Mechanical Supply – {total_supply} m3/s')
                        # check if above not zero
                        # Mechanical Supply – 3.3 m3/s
                        # 1.5m2 AOV; 1.5m2 Natural Smoke Shaft, 1.0m2 Corridor Vent; Though Stair Door via 1.0m2 AOV


                        aov = venting_obj["stair_aov"]["area"]
                        if aov:
                            cell_text_list.append(f'{aov} m2 AOV') # how to know if aov provided for head of stair or corridor?

                        natural_inlet = sum(venting_obj["natural_openings"]) # check with Sam with artificial/for modelling inlet to be included?
                        natural_inlet = round_to(natural_inlet)
                        # if natural_inlet:
                        #     cell_text_list.append(f'Additional inlet modelled in room of origin – {natural_inlet} m2') # how to superscript??

                        cell_text = "; ".join(cell_text_list)

                        replace_table_cell_content(cell=row_cells[column_index], replacement_text=f'ENGINEER TO CONFIRM: {cell_text}')
                        # how to get type of supply??
                # column 4 inlet type

            for index_scenario in range(len(scenario_names)):
                scenario_table_row(table_object=scenario_table, index_scenario=index_scenario)

            def locate_table_from_cell(cell_row_index, cell_column_index, cell_text):
                table_list = []
                
                for table in document.tables:
                    # check table has enough rows and columns
                    cell = table.rows[cell_row_index].cells[cell_column_index]
                    if cell_text in cell.text:
                        table_list.append(table)
                return table_list
                # return all occurrences
            def return_general_scen(firefighting=False):
                if firefighting:
                    return scenarios_object[FSA_scenarios[0]]
                else:
                    return scenarios_object[MoE_scenarios[0]]
            def scenario_timeline_table(table_object, firefighting=False):
                table_rows = table_object.rows
                # loop through rows
                # if ff = False use first scenario for opening times
                # else use last
                scen_for_timings = return_general_scen(firefighting)

                for row_index in range(len(table_rows)):
                    # first cell text
                    row_cells = table_rows[row_index].cells
                    row_title_cell = row_cells[0]
                    if "Apartment Door Open" in row_title_cell.text:
                        # column index 1 => stair door opening - add to opening object
                        target_cell = table_rows[row_index].cells[1]
                        cell_text = str(scen_for_timings["door_opening_times"]["opening_apartment"])
                        replace_table_cell_content(cell=target_cell, replacement_text=cell_text)
                    if "Stair Door Open" in row_title_cell.text:
                        # column index 1 => stair door opening - add to opening object
                        target_cell = table_rows[row_index].cells[1]
                        # TODO: needs to be rounded
                        cell_text = str(scen_for_timings["door_opening_times"]["opening_stair"])
                        replace_table_cell_content(cell=target_cell, replacement_text=cell_text)
                    if "Apartment Door Close" in row_title_cell.text:
                        # column index 1 => stair door opening - add to opening object
                        target_cell = table_rows[row_index].cells[1]
                        cell_text = str(scen_for_timings["door_opening_times"]["closing_apartment"])
                        replace_table_cell_content(cell=target_cell, replacement_text=cell_text)
                        #  likely not needed column index 3 => scope whether stair door closes or not
                        # probably leave max TD  and xxs based on walking speed of 1.2m/s
                    if "Stair Door Close" in row_title_cell.text:
                        target_cell = table_rows[row_index].cells[1]
                        cell_text = str(scen_for_timings["door_opening_times"]["closing_stair"])
                        replace_table_cell_content(cell=target_cell, replacement_text=cell_text)
                        # if row_index == 7:
                        # stair door closing -> add to opening object
                        # model t
                    if "Terminate" in row_title_cell.text:
                        target_cell = table_rows[row_index].cells[1]
                        cell_text = str(scen_for_timings["end_time"])
                        replace_table_cell_content(cell=target_cell, replacement_text=cell_text)
                        # model max t
                # first change @ row index 4

                # take timings from object
                # need maxTime in model
                        
            MoE_FSA_tables = locate_table_from_cell(0,0,"Event")
            # Jinja removes timeline tables and sections when no FSA or no MOE present

            def populate_results_section(firefighting=False):
        #     #     # bullet points where only 1 scen for fsa or MOE in section above
                        #     #     # using jinja 

                    results_tables = locate_table_from_cell(0,-1,"Meets Performance")
                    if not firefighting: # i.e. MoE
                        # if > 1 in scenarios
                        moe_results_table = results_tables[0] 
                        alter_table_rows(total_rows=len(MoE_scenarios), table=moe_results_table)
                        table_rows = moe_results_table.rows
                        for index in range(len(MoE_scenarios)):
                        # use index for row index
                            row_index = index + 1 # top row for headings
                            row_cells = table_rows[row_index].cells
                            

                            scenario = MoE_scenarios[index]
                            # obtain from helper
                            tenable_time, max_pressure_drop, meet_criteria = scen_results_values(scenario, scenarios_object, firefighting)
                            # for bullets set varible to above values
                            replace_table_cell_content(cell=row_cells[1], replacement_text=str(round_to(tenable_time)))

                            replace_table_cell_content(cell=row_cells[2], replacement_text=f'{max_pressure_drop}kPa')
                            replace_table_cell_content(cell=row_cells[3], replacement_text=meet_criteria)
                    elif firefighting:
                        fsa_results_table = results_tables[-1] # if > 1 in scenarios; else bullet
                        alter_table_rows(total_rows=len(FSA_scenarios), table=fsa_results_table, header_rows=2)
                        table_rows = fsa_results_table.rows                   
                        #     if len(fsa_or_moe_scenarios) == 1:
                        for index in range(len(FSA_scenarios)):
                            # use index for row index
                                row_index = index + 2 # top 2 row for headings; double row
                                row_cells = table_rows[row_index].cells
                                

                                scenario = FSA_scenarios[index]
                                tenable_object = scenarios_object[scenario]["tenability"] # applies to bullets
                                tenability_keys = list(tenable_object.keys())   # this and below applies to bullets
                                text_list, worst_temp, worst_vis, meet_criteria, max_pressure_drop = scen_results_values(scenario, scenarios_object, firefighting)
                                for index_key in range(len(text_list)):
                                    replace_table_cell_content(cell=row_cells[index_key+1], replacement_text=(text_list[index_key]))

                                replace_table_cell_content(cell=row_cells[4], replacement_text=str(round_to(worst_vis)))
                                replace_table_cell_content(cell=row_cells[5], replacement_text=str(round_to(worst_temp)))
                                replace_table_cell_content(cell=row_cells[-2], replacement_text=str((max_pressure_drop)))
                                replace_table_cell_content(cell=row_cells[-1], replacement_text=meet_criteria)

                
            # if MoE >0; then will be first table
            if len(MoE_scenarios) > 0:
                moe_table = MoE_FSA_tables[0]
                scenario_timeline_table(table_object=moe_table, firefighting=False)            
                if len(MoE_scenarios) > 1: # if moe > 1; else would be bullet points      
                    populate_results_section(firefighting=False)
            if len(FSA_scenarios) > 0:
                fsa_table = MoE_FSA_tables[-1] 
                scenario_timeline_table(table_object=fsa_table, firefighting=True)
                if len(FSA_scenarios) > 1: # if > 1 use table for results, else populate bullets
                    populate_results_section(firefighting=True)

            document.save(output_path)
            startup_path = rf'{os.path.abspath(os.getcwd())}/{output_path}'
            os.startfile(startup_path)

            window.close()


    # Below for without gui

    # {{PROJECT NAME}} {{TODAYS DATE}} {{CLIENT NAME}}

    # project_name = "Test Resi Project"
    # client_name = "Test Client"

    # context = {
    #     "PROJECT_NAME": project_name,
    #     "TODAYS_DATE": today.strftime("%d-%m-%Y"),
    #     "CLIENT_NAME": client_name
    # }
    # doc.render(context)
    # doc.save(Path(__file__).parent / f"{project_name}-Fire Dynamics.docx")