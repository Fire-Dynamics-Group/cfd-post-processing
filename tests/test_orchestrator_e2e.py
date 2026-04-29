"""End-to-end test for the real orchestrator against the Finchley run dir.

This is the spec for PR 2 step 6: invoking ``run_orchestrator`` on a real
FDS run directory must produce a valid, populated ``.docx`` and emit each
phase in order — without GUI popups, without the legacy ``is_test`` Easter
egg, and without ``os.startfile`` auto-opening Word.

Marked ``slow`` because it generates real charts (matplotlib) and renders
a real docx; default ``pytest`` runs (``-m "not slow"``) skip it.

Run explicitly:  pytest -m slow tests/test_orchestrator_e2e.py
"""
from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

from pipeline.services.job import Step
from pipeline.services.report import ReportRequest, run_orchestrator


FINCHLEY_RUN_DIR = Path(
    r"C:\Users\IanShaw\Fire Dynamics Group Dropbox"
    r"\03 Modelling Data"
    r"\0406 - Former Homebase Site North Finchley"
    r"\Rev00 Models"
)

needs_finchley = pytest.mark.skipif(
    not FINCHLEY_RUN_DIR.exists(),
    reason="Finchley fixture not on this machine",
)


@needs_finchley
@pytest.mark.slow
def test_orchestrator_produces_valid_docx_against_finchley(tmp_path: Path) -> None:
    seen_steps: list[Step] = []

    def on_progress(step: Step, _pct: float) -> None:
        if not seen_steps or seen_steps[-1] is not step:
            seen_steps.append(step)

    req = ReportRequest(
        PATH=str(FINCHLEY_RUN_DIR),
        CLIENT_NAME="Test Client",
        PROJECT_NAME="Finchley_E2E",
        PROJECT_LOCATION="North Finchley",
        EMAIL_PREFIX="ian",
        HAS_EXTENDED_TRAVEL=True,
        MAX_TD=15,
        GUIDANCE="BS9991",
        OUTPUT_DIR=str(tmp_path),
    )

    # Belt-and-braces guard: the legacy code called ``os.startfile`` to
    # auto-open Word. The orchestrator MUST NOT do that — the frontend
    # exposes an explicit Open in Word button instead.
    with patch("os.startfile", create=True) as startfile_mock:
        result = run_orchestrator(req, on_progress)

    assert startfile_mock.call_count == 0, "orchestrator must not auto-open Word"

    # All five real phases emitted in order.
    assert seen_steps == [
        Step.PARSING,
        Step.CHARTING,
        Step.DRAWING,
        Step.RENDERING,
        Step.SAVING,
    ]

    # Output is a real, readable .docx in OUTPUT_DIR.
    out = Path(result.output_path)
    assert out.exists(), f"orchestrator did not write {out}"
    assert out.parent == tmp_path
    assert out.name.startswith(f"{req.PROJECT_NAME}-CFD Report-")
    assert out.name.endswith(".docx")

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    # Sanity: project name and client name surface in the rendered doc.
    assert req.PROJECT_NAME in full_text or any(
        req.PROJECT_NAME in cell.text
        for tbl in doc.tables for row in tbl.rows for cell in row.cells
    )

    # The PR 2 part-A placeholder marker must be GONE.
    assert "PLACEHOLDER" not in full_text.upper()
    assert "scaffold" not in full_text.lower()

    # The doc must contain something — empty docs from a failed render
    # would slip through the existence check.
    assert len(doc.paragraphs) > 5
    assert len(doc.tables) > 0
